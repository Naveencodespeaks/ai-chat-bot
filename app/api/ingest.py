from __future__ import annotations

from typing import List, Optional, Dict, Any
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, BackgroundTasks
from sqlalchemy.orm import Session
from datetime import datetime
import json
from app.db.vector import get_qdrant_client


from app.core.logging import get_logger
from app.db.deps import get_db
from app.auth.dependencies import get_current_user
from app.models.user import User
from app.models.documents import Document
from app.schemas.document import DocumentResponse, DocumentCreateRequest
from app.core.config import settings
from app.rag.ingest import ingest_document
from app.auth.context import UserContext

logger = get_logger(__name__)

router = APIRouter(prefix="/api/documents", tags=["documents"])


# -----------------------------
# TEXT EXTRACTION (PRODUCTION)
# -----------------------------
def extract_text_from_upload(file: UploadFile, content: bytes) -> str:
    """
    Extract text from supported file formats.
    """
    filename = (file.filename or "").lower()
    content_type = (file.content_type or "").lower()

    # TXT / MD
    if filename.endswith(".txt") or filename.endswith(".md") or content_type.startswith("text/"):
        return content.decode("utf-8", errors="ignore")

    # PDF
    if filename.endswith(".pdf") or "pdf" in content_type:
        try:
            from pypdf import PdfReader  # pip install pypdf
            import io
            reader = PdfReader(io.BytesIO(content))
            pages = []
            for p in reader.pages:
                pages.append(p.extract_text() or "")
            return "\n\n".join(pages).strip()
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to extract PDF text: {e}",
            )

    # DOCX
    if filename.endswith(".docx") or "word" in content_type:
        try:
            from docx import Document as DocxDocument  # python-docx
            import io
            doc = DocxDocument(io.BytesIO(content))
            return "\n".join([p.text for p in doc.paragraphs]).strip()
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to extract DOCX text: {e}",
            )

    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail="Unsupported file type. Supported: PDF, TXT, MD, DOCX",
    )


def build_user_context(current_user: User) -> UserContext:
    """
    Build UserContext for RBAC-safe ingestion.
    Adjust fields to match your actual UserContext model.
    """
    # If your UserContext has a constructor different from below, tell me your UserContext class.
    return UserContext(
        user_id=current_user.id,
        roles=getattr(current_user, "roles", []) or [],
        department=getattr(current_user, "department", None),
        allowed_visibility=["INTERNAL"],
        is_verified=True,
    )


def parse_metadata(metadata: Optional[str]) -> Dict[str, Any]:
    """
    metadata comes as optional JSON string in multipart.
    """
    if not metadata:
        return {}
    try:
        if isinstance(metadata, str):
            return json.loads(metadata)
        return dict(metadata)
    except Exception:
        return {}


# -----------------------------
# BACKGROUND INGEST TASK
# -----------------------------
def ingest_in_background(
    *,
    doc_id: str,
    text: str,
    title: str,
    source: str,
    meta: Dict[str, Any],
    user_context: UserContext,
):
    """
    Runs after response. Keeps ingestion reliable + non-blocking.
    """
    try:
        from app.db.vector import get_qdrant_client  # ensure you have this function
        client = get_qdrant_client()

        department = meta.get("department")
        allowed_roles = meta.get("allowed_roles")
        visibility = meta.get("visibility")
        tags = meta.get("tags")

        result = ingest_document(
            client=client,
            user_context=user_context,
            text=text,
            title=title,
            source=source,
            department=department,
            allowed_roles=allowed_roles,
            visibility=visibility,
            tags=tags,
            doc_id=doc_id,  # IMPORTANT: ties vectors to Document row id
            db=None,  # keep SQL persistence OFF from this background task unless you want it
        )

        logger.info(f"[RAG] Ingest done doc_id={result.doc_id} chunks={result.chunks_ingested}")

    except Exception as e:
        logger.error(f"[RAG] Ingest failed doc_id={doc_id}: {e}")


# -----------------------------
# ROUTES
# -----------------------------
@router.post(
    "/upload",
    response_model=DocumentResponse,
    status_code=status.HTTP_201_CREATED,
    summary="Upload document",
    description="Upload and ingest a document for RAG",
)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    metadata: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    try:
        allowed_extensions = {".pdf", ".txt", ".md", ".docx"}
        if not file.filename or "." not in file.filename:
            raise HTTPException(status_code=400, detail="Invalid filename")

        file_ext = f".{file.filename.split('.')[-1].lower()}"
        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File type not allowed. Supported: {', '.join(sorted(allowed_extensions))}",
            )

        content = await file.read()
        text_content = extract_text_from_upload(file, content)
        if not text_content.strip():
            raise HTTPException(status_code=400, detail="No text extracted from document")

        meta = parse_metadata(metadata)

        # Create document record (DB source-of-truth)
        document = Document(
            id=f"doc_{int(datetime.now().timestamp() * 1000)}",
            title=file.filename,
            name=file.filename,
            file_path=f"uploads/{file.filename}",
            document_type=file.content_type or "application/octet-stream",
            content_type=file.content_type or "application/octet-stream",
            file_size=len(content),
            owner_id=current_user.id,
            meta=meta or {},
        )

        db.add(document)
        db.commit()
        db.refresh(document)

        user_context = build_user_context(current_user)

        # Background ingest (safe)
        background_tasks.add_task(
            ingest_in_background,
            doc_id=document.id,
            text=text_content,
            title=document.title,
            source=document.file_path,
            meta=meta,
            user_context=user_context,
        )

        logger.info(f"Document uploaded: {document.id} ({file.filename}) user={current_user.id}")
        return DocumentResponse(**document.to_dict())

    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"Error uploading document: {e}")
        raise HTTPException(status_code=500, detail="Failed to upload document")


@router.post(
    "",
    response_model=DocumentResponse,
    status_code=status.HTTP_201_CREATED,
    summary="Create document from content",
    description="Create a document from raw content and ingest",
)
def create_document(
    request: DocumentCreateRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    try:
        meta = request.metadata or {}

        document = Document(
            id=f"doc_{int(datetime.now().timestamp() * 1000)}",
            title=request.title,
            name=request.title,
            file_path=f"content/{request.title}",
            document_type="text",
            content_type="text/plain",
            file_size=len(request.content),
            owner_id=current_user.id,
            meta=meta,
        )

        db.add(document)
        db.commit()
        db.refresh(document)

        user_context = build_user_context(current_user)

        background_tasks.add_task(
            ingest_in_background,
            doc_id=document.id,
            text=request.content,
            title=document.title,
            source=document.file_path,
            meta=meta,
            user_context=user_context,
        )

        logger.info(f"Document created: {document.id} title={request.title} user={current_user.id}")
        return DocumentResponse(**document.to_dict())

    except Exception as e:
        db.rollback()
        logger.error(f"Error creating document: {e}")
        raise HTTPException(status_code=500, detail="Failed to create document")


@router.get(
    "",
    response_model=List[DocumentResponse],
    summary="List documents",
    description="Retrieve all ingested documents",
)
def list_documents(
    skip: int = 0,
    limit: int = 50,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    try:
        documents = db.query(Document).offset(skip).limit(limit).all()
        return [DocumentResponse(**d.to_dict()) for d in documents]
    except Exception as e:
        logger.error(f"Error retrieving documents: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve documents")


@router.get(
    "/{document_id}",
    response_model=DocumentResponse,
    summary="Get document details",
    description="Retrieve details of a specific document",
)
def get_document(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return DocumentResponse(**document.to_dict())

@router.get(
    "/{document_id}/chunks",
    response_model=List[dict],
    summary="Get document chunks",
    description="Retrieve vector store chunks for a document",
)
def get_document_chunks(
    document_id: str,
    limit: int = 20,
    current_user: User = Depends(get_current_user),
):
    """
    Reads from Qdrant with RBAC filter.
    """

    try:
        client = get_qdrant_client()
        from qdrant_client.http import models as qmodels

        user_roles = [r.upper() for r in getattr(current_user, "roles", [])]
        user_department = getattr(current_user, "department", None)

        # -----------------------------
        # ADMIN override
        # -----------------------------
        if "ADMIN" in user_roles or "SUPERADMIN" in user_roles:
            scroll_filter = qmodels.Filter(
                must=[
                    qmodels.FieldCondition(
                        key="document_id",
                        match=qmodels.MatchValue(value=document_id),
                    )
                ]
            )
        else:
            conditions = [
                qmodels.FieldCondition(
                    key="document_id",
                    match=qmodels.MatchValue(value=document_id),
                )
            ]

            if user_department:
                conditions.append(
                    qmodels.FieldCondition(
                        key="department",
                        match=qmodels.MatchValue(value=user_department),
                    )
                )

            if user_roles:
                conditions.append(
                    qmodels.FieldCondition(
                        key="allowed_roles",
                        match=qmodels.MatchAny(any=user_roles),
                    )
                )

            scroll_filter = qmodels.Filter(must=conditions)

        points, _ = client.scroll(
            collection_name=settings.RAG_COLLECTION,
            scroll_filter=scroll_filter,
            limit=limit,
            with_payload=True,
            with_vectors=False,
        )

        return [p.payload for p in points]

    except Exception as e:
        logger.error(f"Error retrieving chunks for document {document_id}: {e}")
        raise HTTPException(
            status_code=500,
            detail="Failed to retrieve document chunks",
        )
